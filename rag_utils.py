from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator

from bs4 import BeautifulSoup
from docx import Document
from openai import OpenAI
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


@dataclass(frozen=True)
class TextUnit:
    text: str
    page: int | None = None


def get_lmstudio_client() -> OpenAI:
    return OpenAI(
        base_url=(os.getenv("LMSTUDIO_BASE_URL") or "http://localhost:1234/v1").strip(),
        api_key=(os.getenv("LMSTUDIO_API_KEY") or "lm-studio").strip() or "lm-studio",
        timeout=120.0,
        max_retries=2,
    )


def prefixed_text(text: str, prefix: str) -> str:
    prefix = prefix.strip()
    return f"{prefix} {text}" if prefix else text


def iter_supported_files(root: Path) -> Iterator[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path


def load_text_units(path: Path) -> list[TextUnit]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    if suffix in {".txt", ".md"}:
        return [TextUnit(_read_text_file(path))]
    raise ValueError(f"Unsupported file type: {suffix}")


def _load_pdf(path: Path) -> list[TextUnit]:
    reader = PdfReader(str(path))
    units: list[TextUnit] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = normalize_text(text)
        if text:
            units.append(TextUnit(text=text, page=page_number))
    return units


def _load_docx(path: Path) -> list[TextUnit]:
    document = Document(str(path))
    blocks: list[str] = []
    blocks.extend(p.text for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                blocks.append(line)
    return [TextUnit(normalize_text("\n\n".join(blocks)))]


def _load_html(path: Path) -> list[TextUnit]:
    raw = _read_text_file(path)
    soup = BeautifulSoup(raw, "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    return [TextUnit(normalize_text(soup.get_text("\n")))]


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return normalize_text(path.read_text(encoding=encoding))
        except UnicodeDecodeError:
            continue
    return normalize_text(path.read_text(encoding="utf-8", errors="replace"))


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, max_chars: int = 1800, overlap_chars: int = 250) -> list[str]:
    """Split text near paragraph/sentence boundaries with a small overlap."""
    text = normalize_text(text)
    if not text:
        return []
    if overlap_chars >= max_chars:
        raise ValueError("overlap_chars must be smaller than max_chars")

    chunks: list[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        hard_end = min(start + max_chars, text_len)
        end = hard_end

        if hard_end < text_len:
            search_start = start + max_chars // 2
            window = text[search_start:hard_end]
            candidates = [
                window.rfind("\n\n"),
                window.rfind(". "),
                window.rfind("? "),
                window.rfind("! "),
                window.rfind("; "),
                window.rfind("\n"),
                window.rfind(" "),
            ]
            boundary = max(candidates)
            if boundary >= 0:
                end = search_start + boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_len:
            break
        start = max(end - overlap_chars, start + 1)

    return chunks


def batched(items: list[str], batch_size: int) -> Iterable[list[str]]:
    for index in range(0, len(items), batch_size):
        yield items[index : index + batch_size]


def embed_texts(client: OpenAI, model: str, texts: list[str], batch_size: int = 24) -> list[list[float]]:
    vectors: list[list[float]] = []
    for batch in batched(texts, batch_size):
        response = client.embeddings.create(model=model, input=batch)
        ordered = sorted(response.data, key=lambda item: item.index)
        vectors.extend(item.embedding for item in ordered)
    return vectors


def stable_chunk_id(source: str, page: int | None, chunk_index: int, text: str) -> str:
    raw = f"{source}|{page}|{chunk_index}|{text}".encode("utf-8", errors="ignore")
    return hashlib.sha1(raw).hexdigest()
